"""
Google Drive AI Agent: Reading & Parsing Files
This extension adds document reading and analysis capabilities to the Google Drive Agent.
"""

import os
import re
import io
import base64
from typing import List, Dict, Any, Optional, Tuple
from pydantic import BaseModel, Field
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
import pickle
import datetime

# Document processing imports
from pypdf import PdfReader
from docx import Document
import pandas as pd
import zipfile
import csv
import json
from bs4 import BeautifulSoup
import markdown
import html2text

# For text analysis
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
import dateutil.parser

from langchain.tools import BaseTool
from langchain.agents import AgentExecutor, create_openai_tools_agent
from langchain_openai import ChatOpenAI
from langchain.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain.memory import ConversationBufferMemory
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains.summarize import load_summarize_chain
from langchain.docstore.document import Document as LangchainDocument
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.chains import RetrievalQA

from app.core.auth import get_drive_service #NOTE: uncomment and use this instead

# Download NLTK resources
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')
try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')


# Define schemas for each tool
class ReadFileInput(BaseModel):
    file_id: str = Field(..., description="The ID of the file to read")
    max_pages: int = Field(default=5, description="Maximum number of pages to read (for PDFs)")

class ParseDocumentInput(BaseModel):
    file_id: str = Field(..., description="The ID of the file to parse")
    parse_level: str = Field(
        default="sections",
        description="Level of parsing: 'sections', 'paragraphs', or 'sentences'",
    )

class ExtractInfoInput(BaseModel):
    file_id: str = Field(..., description="The ID of the file to extract information from")
    info_types: str = Field(
        default="all",
        description="Types of information to extract (comma-separated): 'dates', 'names', 'emails', 'urls', 'headers', 'all'",
    )

class SummarizeDocumentInput(BaseModel):
    file_id: str = Field(..., description="The ID of the file to summarize")
    summary_length: str = Field(
        default="medium", description="Length of summary: 'short', 'medium', or 'long'"
    )

class SearchInDocumentInput(BaseModel):
    file_id: str = Field(..., description="The ID of the file to search in")
    query: str = Field(..., description="The keyword or phrase to search for")
    case_sensitive: bool = Field(default=False, description="Whether the search should be case-sensitive")

class AnswerQuestionInput(BaseModel):
    file_id: str = Field(..., description="The ID of the file to query")
    question: str = Field(..., description="The question to answer based on the file contents")

class FileReader:
    """Class to handle reading different file types from Google Drive."""

    def __init__(self, service):
        self.service = service

    def get_file_metadata(self, file_id):
        """Get metadata for a file."""
        return self.service.files().get(fileId=file_id, fields="name, mimeType").execute()

    def download_file(self, file_id):
        """Download a file's content."""
        request = self.service.files().get_media(fileId=file_id)
        file_content = io.BytesIO()
        downloader = MediaIoBaseDownload(file_content, request)
        done = False
        while not done:
            _, done = downloader.next_chunk()
        file_content.seek(0)
        return file_content

    def export_google_doc(self, file_id, mime_type='text/plain'):
        """Export a Google Doc to the specified format."""
        request = self.service.files().export_media(fileId=file_id, mimeType=mime_type)
        file_content = io.BytesIO()
        downloader = MediaIoBaseDownload(file_content, request)
        done = False
        while not done:
            _, done = downloader.next_chunk()
        file_content.seek(0)
        return file_content

    def read_file(self, file_id, max_pages=5):
        """Read a file's content based on its type."""
        try:
            file_metadata = self.get_file_metadata(file_id)
            file_name = file_metadata.get('name', 'Unknown')
            mime_type = file_metadata.get('mimeType', 'Unknown')

            # Handle Google Docs
            if mime_type == "application/vnd.google-apps.document":
                content = self.export_google_doc(file_id).read().decode("utf-8")
                return {
                    "content": content,
                    "file_name": file_name,
                    "mime_type": mime_type,
                    "status": "success",
                }

            # Handle Google Sheets
            elif mime_type == 'application/vnd.google-apps.spreadsheet':
                content = self.export_google_doc(file_id, mime_type='text/csv').read().decode('utf-8')
                return {
                    'content': content,
                    'file_name': file_name,
                    'mime_type': mime_type,
                    'status': 'success'
                }

            # Handle Google Slides
            elif mime_type == "application/vnd.google-apps.presentation":
                content = (
                    self.export_google_doc(file_id, mime_type="text/plain")
                    .read()
                    .decode("utf-8")
                )
                return {
                    "content": content,
                    "file_name": file_name,
                    "mime_type": mime_type,
                    "status": "success",
                }

            # Handle PDFs
            elif mime_type == 'application/pdf':
                file_content = self.download_file(file_id)
                pdf_reader = PdfReader(file_content)
                content = ""
                for i in range(min(len(pdf_reader.pages), max_pages)):
                    content += pdf_reader.pages[i].extract_text() + "\n\n"

                if len(pdf_reader.pages) > max_pages:
                    content += f"\n[Note: Only showing first {max_pages} of {len(pdf_reader.pages)} pages]"

                return {
                    'content': content,
                    'file_name': file_name,
                    'mime_type': mime_type,
                    'status': 'success',
                    'pages': len(pdf_reader.pages),
                    'pages_read': min(len(pdf_reader.pages), max_pages)
                }

            # Handle DOCX
            elif (
                mime_type
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                file_content = self.download_file(file_id)
                doc = Document(file_content)
                content = "\n".join([para.text for para in doc.paragraphs])
                return {
                    "content": content,
                    "file_name": file_name,
                    "mime_type": mime_type,
                    "status": "success",
                }

            # Handle Plain Text
            elif mime_type in ['text/plain', 'text/markdown', 'text/csv', 'application/json']:
                file_content = self.download_file(file_id)
                content = file_content.read().decode('utf-8')
                return {
                    'content': content,
                    'file_name': file_name,
                    'mime_type': mime_type,
                    'status': 'success'
                }

            # Handle HTML
            elif mime_type in ["text/html"]:
                file_content = self.download_file(file_id)
                html_content = file_content.read().decode("utf-8")
                h = html2text.HTML2Text()
                h.ignore_links = False
                content = h.handle(html_content)
                return {
                    "content": content,
                    "file_name": file_name,
                    "mime_type": mime_type,
                    "status": "success",
                }

            # Unsupported type
            else:
                return {
                    'content': None,
                    'file_name': file_name,
                    'mime_type': mime_type,
                    'status': 'error',
                    'error': f"Unsupported file type: {mime_type}"
                }

        except Exception as e:
            return {
                "content": None,
                "file_name": "Unknown",
                "mime_type": "Unknown",
                "status": "error",
                "error": str(e),
            }


class DocumentParser:
    """Class to parse documents into sections, paragraphs, or sentences."""

    @staticmethod
    def parse_document(content, level='sections'):
        """Parse a document into the specified level of granularity."""
        if not content:
            return []

        if level == "sections":
            # Define section patterns based on headings
            section_patterns = [
                # Markdown/HTML headings
                r"(?:^|\n)#{1,6}\s+(.+?)(?=\n#{1,6}\s+|\Z)",
                # Numbered headings
                r"(?:^|\n)(?:\d+\.)+\s+(.+?)(?=\n(?:\d+\.)+\s+|\Z)",
                # Uppercase headings
                r"(?:^|\n)([A-Z][A-Z\s]+[A-Z])(?:\n|\Z)",
                # Lines ending with colon (potential section headers)
                r"(?:^|\n)([^:\n]+):(?:\n|\Z)",
            ]

            sections = []
            remaining_content = content

            for pattern in section_patterns:
                matches = re.finditer(pattern, content, re.MULTILINE | re.DOTALL)
                for match in matches:
                    section_title = match.group(1).strip()
                    section_content = match.group(0).strip()
                    sections.append(
                        {"title": section_title, "content": section_content}
                    )
                    # Remove matched content from remaining_content
                    remaining_content = remaining_content.replace(match.group(0), "", 1)

            # If no sections were found or there's remaining content, create default sections
            if not sections or remaining_content.strip():
                # Try to split by double newlines
                parts = [p for p in remaining_content.split('\n\n') if p.strip()]
                if len(parts) > 1:
                    for i, part in enumerate(parts):
                        sections.append({
                            'title': f"Section {i+1}",
                            'content': part.strip()
                        })
                else:
                    sections.append({
                        'title': "Document Content",
                        'content': remaining_content.strip()
                    })

            return sections

        elif level == 'paragraphs':
            # Split by double newlines and filter out empty paragraphs
            paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
            return [{'paragraph': i+1, 'content': p} for i, p in enumerate(paragraphs)]

        elif level == "sentences":
            # Use NLTK to tokenize into sentences
            sentences = sent_tokenize(content)
            return [
                {"sentence": i + 1, "content": s.strip()}
                for i, s in enumerate(sentences)
                if s.strip()
            ]

        else:
            # Default to returning the raw content
            return [{'content': content}]

class InformationExtractor:
    """Class to extract various types of information from text."""

    @staticmethod
    def extract_dates(text):
        """Extract dates from text."""
        # Common date patterns
        date_patterns = [
            r"\b(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{1,2}(?:st|nd|rd|th)?,?\s+\d{4}\b",
            r"\b\d{1,2}(?:st|nd|rd|th)?\s+(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?),?\s+\d{4}\b",
            r"\b\d{1,2}/\d{1,2}/\d{2,4}\b",
            r"\b\d{4}-\d{1,2}-\d{1,2}\b",
            r"\b\d{1,2}-\d{1,2}-\d{4}\b",
            r"\b(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{1,2}(?:st|nd|rd|th)?\b",
        ]

        dates = []
        for pattern in date_patterns:
            dates.extend(re.findall(pattern, text, re.IGNORECASE))

        return list(set(dates))  # Remove duplicates

    @staticmethod
    def extract_names(text):
        """Extract potential person names from text."""
        # Simple pattern for potential names (2-3 capitalized words)
        name_patterns = [
            r'\b[A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\b'
        ]

        names = []
        for pattern in name_patterns:
            names.extend(re.findall(pattern, text))

        # Filter out common false positives
        false_positives = {
            "United States",
            "New York",
            "Los Angeles",
            "San Francisco",
            "Hong Kong",
            "United Kingdom",
            "Monday Morning",
            "Tuesday Evening",
            "Wednesday Afternoon",
            "Thursday Night",
            "Friday Morning",
            "Saturday Evening",
            "Sunday Afternoon",
        }

        filtered_names = [name for name in names if name not in false_positives]
        return list(set(filtered_names))  # Remove duplicates

    @staticmethod
    def extract_emails(text):
        """Extract email addresses from text."""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        return list(set(re.findall(email_pattern, text)))

    @staticmethod
    def extract_urls(text):
        """Extract URLs from text."""
        url_pattern = r"https?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+[/\w\.-]*\b"
        return list(set(re.findall(url_pattern, text)))

    @staticmethod
    def extract_headers(text):
        """Extract potential headers from text."""
        # Patterns for different header styles
        header_patterns = [
            # Markdown headers
            r'(?:^|\n)#{1,6}\s+(.+?)(?=\n|$)',
            # Uppercase "headers"
            r'(?:^|\n)([A-Z][A-Z\s]+[A-Z])(?:\n|$)',
            # Numbered headers
            r'(?:^|\n)(?:\d+\.)+\s+(.+?)(?=\n|$)',
            # Underlined headers with = or -
            r'(?:^|\n)(.+)\n[=]{2,}(?:\n|$)',
            r'(?:^|\n)(.+)\n[-]{2,}(?:\n|$)'
        ]

        headers = []
        for pattern in header_patterns:
            matches = re.finditer(pattern, text, re.MULTILINE)
            for match in matches:
                if len(match.groups()) > 0:
                    headers.append(match.group(1).strip())
                else:
                    headers.append(match.group(0).strip())

        return list(set(headers))

    @staticmethod
    def extract_information(text, info_types="all"):
        """Extract specified types of information from text."""
        if not text:
            return {}

        requested_types = [t.strip().lower() for t in info_types.split(',')]
        extract_all = 'all' in requested_types

        result = {}

        if extract_all or 'dates' in requested_types:
            result['dates'] = InformationExtractor.extract_dates(text)

        if extract_all or "names" in requested_types:
            result["names"] = InformationExtractor.extract_names(text)

        if extract_all or 'emails' in requested_types:
            result['emails'] = InformationExtractor.extract_emails(text)

        if extract_all or "urls" in requested_types:
            result["urls"] = InformationExtractor.extract_urls(text)

        if extract_all or 'headers' in requested_types:
            result['headers'] = InformationExtractor.extract_headers(text)

        return result


# Tools Implementation
class ReadFileTool(BaseTool):
    name: str = "read_file"
    description: str = "Reads the content of text-based files in Google Drive (Google Docs, TXT, PDFs, etc.). Use this to access the content of a specific file."
    args_schema: type[ReadFileInput] = ReadFileInput
    file_cache: Dict[str, Any] = Field(default_factory=dict)

    def _run(self, file_id: str, max_pages: int = 5) -> str:
        """Reads the content of a file in Google Drive."""
        try:
            service = get_drive_service()
            file_reader = FileReader(service)
            result = file_reader.read_file(file_id, max_pages)

            if result["status"] == "error":
                return f"Error reading file: {result['error']}"

            file_name = result['file_name']
            mime_type = result['mime_type']
            content = result['content']

            # Prepare a summary of what was read
            output = f"Successfully read '{file_name}' ({mime_type}).\n\n"

            # If it's a PDF, add page information
            if 'pages' in result:
                output += f"PDF document with {result['pages']} pages. Read {result['pages_read']} page(s).\n\n"

            # Add a short preview of the content
            preview_length = min(500, len(content))
            output += f"Content preview:\n{content[:preview_length]}"
            if len(content) > preview_length:
                output += "...\n[Content truncated for preview]"

            # Store the full content in the tool's memory for other tools to use
            self.file_cache[file_id] = {
                'file_name': file_name,
                'mime_type': mime_type,
                'content': content,
                'access_time': datetime.datetime.now()
            }

            return output

        except Exception as e:
            return f"Error reading file: {str(e)}"

class ParseDocumentTool(BaseTool):
    name: str = "parse_document"
    description: str = "Parses a document into sections, paragraphs, or sentences for better analysis. Use this to break down document structure."
    args_schema: type[ParseDocumentInput] = ParseDocumentInput
    file_cache: Dict[str, Any] = Field(default_factory=dict)

    def _run(self, file_id: str, parse_level: str = "sections") -> str:
        """Parses a document into the specified level of granularity."""
        try:
            # Check if we have the file content cached
            if hasattr(self, "file_cache") and file_id in self.file_cache:
                file_data = self.file_cache[file_id]
                content = file_data["content"]
                file_name = file_data["file_name"]
            else:
                # Read the file if not cached
                service = get_drive_service()
                file_reader = FileReader(service)
                result = file_reader.read_file(file_id)

                if result['status'] == 'error':
                    return f"Error reading file: {result['error']}"

                content = result["content"]
                file_name = result["file_name"]

                # Cache the content for later use
                self.file_cache[file_id] = {
                    'file_name': result['file_name'],
                    'mime_type': result['mime_type'],
                    'content': content,
                    'access_time': datetime.datetime.now()
                }

            # Parse the document
            parser = DocumentParser()
            parsed_content = parser.parse_document(content, parse_level)

            # Format the output
            if parse_level == 'sections':
                output = f"Parsed '{file_name}' into {len(parsed_content)} sections:\n\n"
                for i, section in enumerate(parsed_content, 1):
                    title = section['title']
                    content_preview = section['content'][:100] + "..." if len(section['content']) > 100 else section['content']
                    output += f"{i}. {title}\n   Preview: {content_preview}\n\n"

            elif parse_level == "paragraphs":
                output = (
                    f"Parsed '{file_name}' into {len(parsed_content)} paragraphs:\n\n"
                )
                for i, para in enumerate(parsed_content[:10], 1):  # Show only first 10
                    content_preview = (
                        para["content"][:100] + "..."
                        if len(para["content"]) > 100
                        else para["content"]
                    )
                    output += f"Paragraph {i}: {content_preview}\n\n"

                if len(parsed_content) > 10:
                    output += f"[{len(parsed_content) - 10} more paragraphs not shown]\n"

            elif parse_level == "sentences":
                output = (
                    f"Parsed '{file_name}' into {len(parsed_content)} sentences:\n\n"
                )
                for i, sentence in enumerate(
                    parsed_content[:15], 1
                ):  # Show only first 15
                    output += f"Sentence {i}: {sentence['content']}\n"

                if len(parsed_content) > 15:
                    output += f"\n[{len(parsed_content) - 15} more sentences not shown]\n"

            else:
                output = f"Unknown parse level: {parse_level}. Please use 'sections', 'paragraphs', or 'sentences'."

            return output

        except Exception as e:
            return f"Error parsing document: {str(e)}"


class ExtractInfoTool(BaseTool):
    name: str = "extract_information"
    description: str = "Extracts key information like dates, names, emails, URLs, and headers from a document. Use this to identify important elements in a file."
    args_schema: type[ExtractInfoInput] = ExtractInfoInput
    file_cache: Dict[str, Any] = Field(default_factory=dict)

    def _run(self, file_id: str, info_types: str = "all") -> str:
        """Extracts key information from a document."""
        try:
            # Check if we have the file content cached
            if hasattr(self, 'file_cache') and file_id in self.file_cache:
                file_data = self.file_cache[file_id]
                content = file_data['content']
                file_name = file_data['file_name']
            else:
                # Read the file if not cached
                service = get_drive_service()
                file_reader = FileReader(service)
                result = file_reader.read_file(file_id)

                if result["status"] == "error":
                    return f"Error reading file: {result['error']}"

                content = result['content']
                file_name = result['file_name']

                # Cache the content for later use
                self.file_cache[file_id] = {
                    "file_name": result["file_name"],
                    "mime_type": result["mime_type"],
                    "content": content,
                    "access_time": datetime.datetime.now(),
                }

            # Extract information
            extractor = InformationExtractor()
            extracted_info = extractor.extract_information(content, info_types)

            # Format the output
            output = f"Information extracted from '{file_name}':\n\n"

            if 'dates' in extracted_info:
                output += "📅 Dates:\n"
                if extracted_info['dates']:
                    for i, date in enumerate(extracted_info['dates'][:15], 1):
                        output += f"  {i}. {date}\n"
                    if len(extracted_info['dates']) > 15:
                        output += f"  [and {len(extracted_info['dates']) - 15} more dates...]\n"
                else:
                    output += "  No dates found.\n"
                output += "\n"

            if "names" in extracted_info:
                output += "👤 Potential Names:\n"
                if extracted_info["names"]:
                    for i, name in enumerate(extracted_info["names"][:15], 1):
                        output += f"  {i}. {name}\n"
                    if len(extracted_info["names"]) > 15:
                        output += f"  [and {len(extracted_info['names']) - 15} more names...]\n"
                else:
                    output += "  No names found.\n"
                output += "\n"

            if 'emails' in extracted_info:
                output += "📧 Email Addresses:\n"
                if extracted_info['emails']:
                    for i, email in enumerate(extracted_info['emails'][:15], 1):
                        output += f"  {i}. {email}\n"
                    if len(extracted_info['emails']) > 15:
                        output += f"  [and {len(extracted_info['emails']) - 15} more emails...]\n"
                else:
                    output += "  No email addresses found.\n"
                output += "\n"

            if "urls" in extracted_info:
                output += "🔗 URLs:\n"
                if extracted_info["urls"]:
                    for i, url in enumerate(extracted_info["urls"][:15], 1):
                        output += f"  {i}. {url}\n"
                    if len(extracted_info["urls"]) > 15:
                        output += (
                            f"  [and {len(extracted_info['urls']) - 15} more URLs...]\n"
                        )
                else:
                    output += "  No URLs found.\n"
                output += "\n"

            if 'headers' in extracted_info:
                output += "📑 Document Headers:\n"
                if extracted_info['headers']:
                    for i, header in enumerate(extracted_info['headers'][:15], 1):
                        output += f"  {i}. {header}\n"
                    if len(extracted_info['headers']) > 15:
                        output += f"  [and {len(extracted_info['headers']) - 15} more headers...]\n"
                else:
                    output += "  No headers found.\n"
                output += "\n"

            return output

        except Exception as e:
            return f"Error extracting information: {str(e)}"

class SummarizeDocumentTool(BaseTool):
    name: str = "summarize_document"
    description: str = "Creates a concise summary of a document. Useful for quickly understanding the main points without reading the entire file."
    args_schema: type[SummarizeDocumentInput] = SummarizeDocumentInput
    file_cache: Dict[str, Any] = Field(default_factory=dict)

    def _run(self, file_id: str, summary_length: str = "medium") -> str:
        """Summarizes a document."""
        try:
            # Check if we have the file content cached
            if hasattr(self, "file_cache") and file_id in self.file_cache:
                file_data = self.file_cache[file_id]
                content = file_data["content"]
                file_name = file_data["file_name"]
            else:
                # Read the file if not cached
                service = get_drive_service()
                file_reader = FileReader(service)
                result = file_reader.read_file(file_id)

                if result['status'] == 'error':
                    return f"Error reading file: {result['error']}"

                content = result["content"]
                file_name = result["file_name"]

                # Cache the content for later use
                self.file_cache[file_id] = {
                    'file_name': result['file_name'],
                    'mime_type': result['mime_type'],
                    'content': content,
                    'access_time': datetime.datetime.now()
                }

            # Determine summary parameters based on requested length
            if summary_length.lower() == "short":
                max_tokens = 150
            elif summary_length.lower() == "long":
                max_tokens = 500
            else:  # medium or default
                max_tokens = 250

            # Create a summary using LangChain's summarization
            llm = ChatOpenAI(temperature=0)
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=2000,
                chunk_overlap=200
            )

            # Split the text into chunks for processing
            texts = text_splitter.split_text(content)

            # Convert to LangChain documents
            docs = [LangchainDocument(page_content=t) for t in texts]

            # Load the summarize chain
            chain = load_summarize_chain(llm, chain_type="map_reduce", verbose=False)

            # Generate summary
            summary = chain.invoke(docs)

            # Format the output
            output = f"Summary of '{file_name}':\n\n{summary}\n\n"
            output += (
                f"Note: This is a {summary_length} summary of the document content."
            )

            return output

        except Exception as e:
            return f"Error summarizing document: {str(e)}"


class SearchInDocumentTool(BaseTool):
    name: str = "search_in_document"
    description: str = "Searches for keywords or phrases within a document. Use this to find specific information in a file."
    args_schema: type[SearchInDocumentInput] = SearchInDocumentInput
    file_cache: Dict[str, Any] = Field(default_factory=dict)

    def _run(self, file_id: str, query: str, case_sensitive: bool = False) -> str:
        """Searches for keywords or phrases within a document."""
        try:
            # Check if we have the file content cached
            if hasattr(self, 'file_cache') and file_id in self.file_cache:
                file_data = self.file_cache[file_id]
                content = file_data['content']
                file_name = file_data['file_name']
            else:
                # Read the file if not cached
                service = get_drive_service()
                file_reader = FileReader(service)
                result = file_reader.read_file(file_id)

                if result["status"] == "error":
                    # Use the specific error message for each tool
                    return f"Error searching document: {result['error']}"  # For SearchInDocumentTool

                content = result["content"]
                file_name = result["file_name"]

                # Cache the content for later use
                self.file_cache[file_id] = {
                    'file_name': result['file_name'],
                    'mime_type': result['mime_type'],
                    'content': content,
                    'access_time': datetime.datetime.now()
                }

            # Split content into sentences for context
            sentences = sent_tokenize(content)

            # Perform the search
            results = []
            search_flags = re.IGNORECASE if not case_sensitive else 0

            for i, sentence in enumerate(sentences):
                if re.search(re.escape(query), sentence, search_flags):
                    # Get some context (previous and next sentence when available)
                    start_idx = max(0, i - 1)
                    end_idx = min(len(sentences) - 1, i + 1)

                    context = " ".join(sentences[start_idx:end_idx + 1])

                    # Highlight the match
                    highlighted = re.sub(
                        f"({re.escape(query)})",
                        r"**\1**",  # Bold in markdown
                        sentence,
                        flags=search_flags,
                    )

                    results.append({
                        'sentence_num': i + 1,
                        'highlighted': highlighted,
                        'context': context
                    })

            # Format the output
            if results:
                output = (
                    f"Found {len(results)} matches for '{query}' in '{file_name}':\n\n"
                )

                for i, result in enumerate(results[:10], 1):
                    output += f"{i}. Match in sentence {result['sentence_num']}:\n"
                    output += f"   {result['highlighted']}\n\n"
                    output += f"   Context: {result['context']}\n\n"

                if len(results) > 10:
                    output += f"[{len(results) - 10} more matches not shown]\n"
            else:
                output = f"No matches found for '{query}' in '{file_name}'."

            return output

        except Exception as e:
            return f"Error searching in document: {str(e)}"


class AnswerQuestionTool(BaseTool):
    name: str = "answer_question"
    description: str = "Answers specific questions about the file contents using RAG (Retrieval-Augmented Generation). Use this to get precise information from the document."
    args_schema: type[AnswerQuestionInput] = AnswerQuestionInput
    file_cache: Dict[str, Any] = Field(default_factory=dict)

    def _run(self, file_id: str, question: str) -> str:
        """Answers a question based on file contents."""
        try:
            # Check if we have the file content cached
            if hasattr(self, 'file_cache') and file_id in self.file_cache:
                file_data = self.file_cache[file_id]
                content = file_data['content']
                file_name = file_data['file_name']
            else:
                # Read the file if not cached
                service = get_drive_service()
                file_reader = FileReader(service)
                result = file_reader.read_file(file_id)

                if result["status"] == "error":
                    return f"Error reading file: {result['error']}"

                content = result['content']
                file_name = result['file_name']

                # Cache the content for later use
                self.file_cache[file_id] = {
                    "file_name": result["file_name"],
                    "mime_type": result["mime_type"],
                    "content": content,
                    "access_time": datetime.datetime.now(),
                }

            # Set up vector store for RAG
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200
            )

            # Split the text into chunks for processing
            texts = text_splitter.split_text(content)

            # Convert to LangChain documents with metadata
            docs = [
                LangchainDocument(
                    page_content=t,
                    metadata={"source": file_name, "file_id": file_id}
                ) for t in texts
            ]

            # Create embeddings
            embeddings = OpenAIEmbeddings()

            # Create vector store
            vectorstore = FAISS.from_documents(docs, embeddings)

            # Create the retrieval QA chain
            qa = RetrievalQA.from_chain_type(
                llm=ChatOpenAI(temperature=0),
                chain_type="stuff",
                retriever=vectorstore.as_retriever(search_kwargs={"k": 3}),
            )

            # Get the answer
            answer = qa.run(question)

            # Format the output
            output = f"Question about '{file_name}':\n"
            output += f"Q: {question}\n\n"
            output += f"A: {answer}"

            return output

        except Exception as e:
            return f"Error answering question: {str(e)}"

